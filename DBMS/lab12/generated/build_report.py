"""Build Lab12_Report.docx mirroring Asad Naeem.docx structure."""
import os
from docx import Document
from docx.shared import Inches, Pt

HERE = os.path.dirname(__file__)
SHOTS = os.path.join(HERE, "shots")

# Section list mirrors the reference 1-to-1.
SECTIONS = [
    ("First installing java and other dependencies:",                       "jver",      6.50),
    ("I also downloaded mysql-connector-j-9.x.x.jar and pasted it into the same folder which would have my java file", None, None),
    ("Now this is my normalized schema for the issue tracker:",             "workbench", 6.50),
    ("Now connecting the database using java :",                            "run",       6.50),
    ("Now adding users:",                                                   "register",  5.10),
    ("Issues:",                                                             "create",    4.45),
    ("Assigning issue to staff:",                                           "assign",    4.75),
    ("Deleting issue (controlled):",                                        "delete",    6.20),
    ("Viewing issues with staff:",                                          "view",      6.50),
    ("Updating status:",                                                    "update",    4.40),
]

SHOT_FILES = {
    "jver":      "shot_jver.png",
    "workbench": "shot_workbench.png",
    "run":       "shot_run.png",
    "register":  "shot_register.png",
    "create":    "shot_create_issue.png",
    "assign":    "shot_assign.png",
    "delete":    "shot_delete.png",
    "view":      "shot_view_join.png",
    "update":    "shot_update.png",
}


def main():
    doc = Document()

    # Title block matches reference: name, lab title, label
    p = doc.add_paragraph()
    r = p.add_run("Arqam Waheed")
    r.font.size = Pt(12)

    doc.add_paragraph("Lab task 12")
    doc.add_paragraph("Data base :")

    for caption, key, width in SECTIONS:
        doc.add_paragraph(caption)
        if key:
            img = os.path.join(SHOTS, SHOT_FILES[key])
            doc.add_picture(img, width=Inches(width))

    out = os.path.join(HERE, "Lab12_Report.docx")
    doc.save(out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
